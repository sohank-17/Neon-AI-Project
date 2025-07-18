from io import BytesIO
from typing import List, Tuple, Union
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from fastapi.responses import StreamingResponse
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from io import BytesIO
import re

def format_messages_for_export(messages: List[dict]) -> str:
    """
    Convert chat messages into a structured exportable string.
    """
    return "\n\n".join([f"{m['role']}:\n{m['content'].strip()}" for m in messages])


def generate_txt_file(text: str) -> BytesIO:
    buffer = BytesIO()
    buffer.write(text.encode("utf-8"))
    buffer.seek(0)
    return buffer


def generate_docx_file(text: str) -> BytesIO:
    doc = Document()
    for block in text.split("\n\n"):
        doc.add_paragraph(block)
        doc.add_paragraph("")  # spacing

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _clean_text_for_pdf(text: str) -> str:
    """
    Clean text for PDF generation to handle special characters and formatting.
    """
    # Remove or replace problematic characters
    text = text.replace('\u2019', "'")  # Smart apostrophe
    text = text.replace('\u2018', "'")  # Smart apostrophe
    text = text.replace('\u201c', '"')  # Smart quote
    text = text.replace('\u201d', '"')  # Smart quote
    text = text.replace('\u2013', '-')  # En dash
    text = text.replace('\u2014', '-')  # Em dash
    
    # Handle markdown-style formatting
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)  # Bold
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)      # Italic
    
    return text


def generate_pdf_file(text: str) -> BytesIO:
    """
    Improved PDF generation with proper text wrapping and formatting.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )

    styles = getSampleStyleSheet()
    
    # Create custom styles for better formatting
    role_style = ParagraphStyle(
        name="RoleStyle",
        parent=styles["Normal"],
        fontSize=12,
        fontName="Helvetica-Bold",
        spaceAfter=6,
        textColor='blue'
    )
    
    content_style = ParagraphStyle(
        name="ContentStyle",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica",
        leading=14,  # Line spacing
        spaceAfter=12,
        leftIndent=20
    )
    
    story = []
    
    # Split text into message blocks
    blocks = text.split("\n\n")
    
    for block in blocks:
        if not block.strip():
            continue
            
        # Clean the text for PDF
        clean_block = _clean_text_for_pdf(block.strip())
        
        # Check if this is a role indicator (user:, assistant:, etc.)
        lines = clean_block.split('\n', 1)
        if len(lines) > 1 and lines[0].strip().endswith(':'):
            # This is a role header
            role = lines[0].strip()
            content = lines[1].strip() if len(lines) > 1 else ""
            
            # Add role header
            story.append(Paragraph(role, role_style))
            
            # Add content if it exists
            if content:
                # Split long content into smaller paragraphs for better formatting
                content_paragraphs = content.split('\n')
                for para in content_paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
        else:
            # Regular content block
            story.append(Paragraph(clean_block, content_style))
        
        # Add some space between message blocks
        story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _render_rich_text(text: str) -> str:
    """
    Convert markdown-style **bold** to reportlab <b>bold</b> tags.
    """
    return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)


def generate_pdf_file_from_blocks(blocks: List[dict]) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)

    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle(
        name="Heading",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        spaceAfter=20
    )

    paragraph_style = ParagraphStyle(
        name="Paragraph",
        parent=styles["Normal"],
        alignment=TA_LEFT,
        leading=14,
        spaceAfter=8
    )

    list_style = ParagraphStyle(
        name="List",
        parent=styles["Normal"],
        leftIndent=20,
        leading=14
    )

    story = []

    for block in blocks:
        if block["type"] == "heading":
            story.append(Paragraph(_render_rich_text(block["text"]), heading_style))
            story.append(Spacer(1, 12))

        elif block["type"] == "paragraph":
            story.append(Paragraph(_render_rich_text(block["text"]), paragraph_style))

        elif block["type"] == "list":
            list_items = [
                ListItem(Paragraph(_render_rich_text(item), list_style))
                for item in block["items"]
            ]
            story.append(
                ListFlowable(
                    list_items,
                    bulletType="1" if block.get("style") == "numbered" else "bullet",
                    start="1" if block.get("style") == "numbered" else None,
                    leftIndent=20
                )
            )
            story.append(Spacer(1, 8))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_chat_as_file(content: Union[str, List[dict]], format: str) -> Tuple[BytesIO, str, str]:
    """
    Export either a list of chat messages or a summary string to the specified format.
    """
    if isinstance(content, list):
        text = format_messages_for_export(content)
    elif isinstance(content, str):
        text = content.strip()
    else:
        raise ValueError("Unsupported content type")

    if format == "txt":
        return generate_txt_file(text), "chat_export.txt", "text/plain"

    elif format == "docx":
        return generate_docx_file(text), "chat_export.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif format == "pdf":
        return generate_pdf_file(text), "chat_export.pdf", "application/pdf"

    else:
        raise ValueError(f"Unsupported export format: {format}")
    

def prepare_export_response(
    content: Union[str, List[dict]],
    format: str,
    filename_prefix: str = "chat_export"
) -> StreamingResponse:
    """
    Prepare a StreamingResponse for export, using the given filename prefix.
    """
    stream, filename, media_type = export_chat_as_file(content, format)

    # Replace "chat_export" with custom prefix if needed
    final_filename = filename.replace("chat_export", filename_prefix)

    return StreamingResponse(
        stream,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={final_filename}"}
    )